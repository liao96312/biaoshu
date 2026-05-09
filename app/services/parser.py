from __future__ import annotations

from pathlib import Path

from openpyxl import load_workbook
from pypdf import PdfReader

from app.adapters.ocr import ocr_client

try:
    from docx import Document as DocxDocument
except Exception:  # pragma: no cover - optional import in minimal environments
    DocxDocument = None


class ParsedDocument:
    def __init__(self, text: str, page_count: int) -> None:
        self.text = text
        self.page_count = page_count


def parse_document(path: str | Path) -> ParsedDocument:
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(file_path)
    if suffix == ".docx":
        return _parse_docx(file_path)
    if suffix in {".xlsx", ".xlsm"}:
        return _parse_xlsx(file_path)
    if suffix in {".txt", ".md"}:
        text = file_path.read_text(encoding="utf-8", errors="ignore")
        return ParsedDocument(text=f"[[page:1]]\n{text}", page_count=1)
    if suffix in {".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff"}:
        text = ocr_client.extract_text(file_path)
        return ParsedDocument(text=f"[[page:1]]\n{text}", page_count=1)
    raise ValueError(f"unsupported file type: {suffix}")


def _parse_pdf(path: Path) -> ParsedDocument:
    reader = PdfReader(str(path))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[[page:{index}]]\n{text}")
    if not pages:
        text = ocr_client.extract_text(path)
        if text.strip():
            pages.append(f"[[page:1]]\n{text}")
    return ParsedDocument(text="\n\n".join(pages), page_count=len(reader.pages))


def _parse_docx(path: Path) -> ParsedDocument:
    if DocxDocument is None:
        raise RuntimeError("python-docx is required to parse .docx files")
    doc = DocxDocument(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    body = "\n".join(paragraphs)
    return ParsedDocument(text=f"[[page:1]]\n{body}", page_count=1)


def _parse_xlsx(path: Path) -> ParsedDocument:
    workbook = load_workbook(path, data_only=True, read_only=True)
    lines: list[str] = []
    try:
        for index, sheet in enumerate(workbook.worksheets, start=1):
            lines.append(f"[[page:{index}]]")
            lines.append(f"[sheet:{sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                values = [str(cell).strip() for cell in row if cell is not None]
                if values:
                    lines.append(" | ".join(values))
        return ParsedDocument(text="\n".join(lines), page_count=len(workbook.worksheets))
    finally:
        workbook.close()


def iter_page_chunks(text: str) -> list[tuple[int | None, str]]:
    chunks: list[tuple[int | None, str]] = []
    current_page: int | None = None
    current_lines: list[str] = []
    for line in text.splitlines():
        line = line.strip()
        if line.startswith("[[page:") and line.endswith("]]"):
            if current_lines:
                chunks.append((current_page, "\n".join(current_lines)))
            raw_page = line.removeprefix("[[page:").removesuffix("]]")
            current_page = int(raw_page) if raw_page.isdigit() else None
            current_lines = []
        elif line:
            current_lines.append(line)
    if current_lines:
        chunks.append((current_page, "\n".join(current_lines)))
    return chunks or [(None, text)]
